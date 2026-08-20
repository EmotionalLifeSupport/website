from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "Emotional-Life-Support-Alternatives-Research-Report-2026-08-15.docx"

FONT = "Calibri"
INK = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5F6B76"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9E0E7"
VERY_LIGHT = "F7F9FB"
GOLD = "9B7B2F"
GREEN = "2F6B4F"
RED = "9B1C1C"
WHITE = "FFFFFF"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM_DXA = 80
CELL_SIDE_DXA = 120


SOURCES = [
    ("Project foundation", "Emotional Life Support Project Foundation and Next-Phase Kickoff Prompt, local working documents, 15 August 2026.", None),
    ("NHS Talking Therapies", "NHS: service scope, self-referral and free access for eligible people in England; page reviewed 5 November 2025.", "https://www.nhs.uk/tests-and-treatments/talking-therapies/"),
    ("BACP therapy access", "British Association for Counselling and Psychotherapy: finding private therapy and advertised fee range.", "https://www.bacp.co.uk/about-therapy/how-to-get-therapy/"),
    ("Counselling Directory price report", "Counselling Directory: January 2026 national average advertised session fee of £60.60, based on its database.", "https://www.counselling-directory.org.uk/press/Counselling-Directory-report/"),
    ("UK life-coach occupation profile", "National Careers Service: life coaching is not regulated in the UK; current occupation guidance.", "https://nationalcareers.service.gov.uk/job-profiles/life-coach"),
    ("ICF credential requirements", "International Coaching Federation: coach-specific education requirements for ACC and PCC credentials.", "https://coachingfederation.org/credentials-and-standards/training-requirements"),
    ("Life-coaching evidence review", "Ammentorp et al. (2013/2014), systematic review: five included studies, limited solid evidence and mixed but promising results.", "https://pmc.ncbi.nlm.nih.gov/articles/PMC4015179/"),
    ("Divorce-intervention meta-analysis", "Strouse and Roehrle (2011), meta-analysis of adult separation/divorce interventions; overall mean effect size 0.47, with heterogeneous interventions.", "https://doi.org/10.1080/14623730.2011.9715665"),
    ("DC² Divorce Coaching", "Public packages: £180 for three sessions, £350 for six, £525 for nine, and £700 for 12 hours of Form E support.", "https://www.dc2divorcecoaching.co.uk/pricing-plans/list"),
    ("Sophie Buck Divorce Coach", "Public packages: £375 monthly; £650 six-week intensive; £995 ten-session programme over three months.", "https://www.sophie-buck.com/pricing-plans/list"),
    ("D Divorce Coaching", "Public packages: £160 single session; £460 for three; £930 for six.", "https://d-divorcecoaching.co.uk/packages/"),
    ("Samantha Johnson Divorce Coach", "Public packages: £520 monthly Bronze; £800 monthly Platinum with weekly sessions, unlimited texts and one emergency call per week.", "https://www.divorcecoachandmarriagetherapy.co.uk/price-packages"),
    ("Emma Rees-Davies Divorce Coach", "Public course ladder: £249 self-guided, £699 group, £899 group plus three 1:1 sessions, £1,099 course plus 12 1:1 sessions.", "https://www.emmareesdaviesdivorcecoaching.co.uk/courses"),
    ("Space For Us", "Breaking Up one-to-one guided course: 16 weekly online sessions, £960 or £800 advance-payment price.", "https://www.spaceforus.co.uk/breaking-up-one-to-one-course"),
    ("Transition Solutions", "Transition-coaching package: consultation, one 120-minute and three 90-minute sessions, £949.", "https://www.transitionsolutions.co.uk/transitioncoaching"),
    ("Nick Hatter", "Public premium coaching packages: £4,320 for six weeks, £7,200 for three months, £12,000 for six months and £18,000 for 12 months.", "https://www.nickhatter.com/prices"),
    ("amicable space", "Divorce/separation membership with weekly content, webinars, forum and consultation: 14-day trial then £20 per month.", "https://amicable.space/become-a-member"),
    ("Divorce Recovery Workshop", "Charity workshops for emotional recovery; 2026 residential examples £190 and £280.", "https://www.drw.org.uk/"),
    ("Divorce court fee", "GOV.UK: £628 divorce application fee in England and Wales, updated 13 July 2026.", "https://www.gov.uk/divorce/file-for-divorce"),
    ("Family mediation", "GOV.UK: usual MIAM cost around £120; voucher of up to £500 for eligible family mediation cases.", "https://www.gov.uk/looking-after-children-divorce/mediation"),
    ("amicable legal services", "Public service prices: £1,200 automated, £1,800 assisted, and £3,450–£5,910 negotiation services.", "https://amicable.io/services/thinking-about-divorce"),
    ("Action PAs", "Public service guide: VA pay-as-you-go £40/hour; five hours £185; 20 hours £740; prices subject to VAT.", "https://action-pas.com/wp-content/uploads/2025/10/Action-PAs-Service-Guide.pdf"),
    ("Transform Each Day", "WhatsApp accountability coaching: £129, £169 and £229 per month with daily check-ins.", "https://www.transformeachday.com/pricing/"),
    ("Accountability Coaching London", "Weekly calls and between-session checkpoints: £590 per month.", "https://accountabilitycoachinglondon.co.uk/full-support-coaching-offer/"),
    ("Wysa App Store listing", "UK App Store prices observed include free access, £11.99 monthly or £66.99 annual Premium, with other in-app purchases.", "https://apps.apple.com/gb/app/wysa-mental-wellbeing-ai/id1166585565"),
    ("Wysa service boundaries", "Wysa FAQ: AI wellbeing support is not a substitute for professional care and is not recommended for crisis situations or severe and enduring mental illness.", "https://www.wysa.com/faq"),
    ("Samaritans", "Free UK phone listening service, 116 123.", "https://www.samaritans.org/how-we-can-help/contact-samaritan/talk-us-phone/"),
    ("Shout 85258", "Free, confidential, 24/7 UK text support delivered by trained volunteers.", "https://giveusashout.org/get-help/"),
    ("Citizens Advice separation guide", "Free guidance on separation choices, mediation, legal help and practical arrangements.", "https://www.citizensadvice.org.uk/family/how-to-separate1/deciding-what-to-do-when-you-separate/"),
]


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run(run, *, size=None, color=None, bold=None, italic=None, font=FONT):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rfonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=CELL_TOP_BOTTOM_DXA, start=CELL_SIDE_DXA, bottom=CELL_TOP_BOTTOM_DXA, end=CELL_SIDE_DXA):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=MID_GRAY, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa=TABLE_INDENT_DXA):
    assert sum(widths_dxa) == CONTENT_WIDTH_DXA, (widths_dxa, sum(widths_dxa))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = tr_pr.find(qn("w:cantSplit"))
        if cant_split is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def keep_lines(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    keep = ppr.find(qn("w:keepLines"))
    if keep is None:
        ppr.append(OxmlElement("w:keepLines"))


def set_paragraph_shading(paragraph, fill: str, border: str | None = None):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    if border:
        pbdr = ppr.find(qn("w:pBdr"))
        if pbdr is None:
            pbdr = OxmlElement("w:pBdr")
            ppr.append(pbdr)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border)
        pbdr.append(left)


def add_hyperlink(paragraph, text: str, url: str, color=BLUE):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([r_fonts, color_node, underline])
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    set_run(run, size=9, color=MUTED)


def ensure_custom_numbering(doc: Document):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(existing_abs, default=0) + 1
    next_num = max(existing_num, default=0) + 1

    def make(kind: str, marker: str, abs_id: int, num_id: int):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abs_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), kind)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), marker)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        ppr.extend([tabs, ind, spacing])
        rpr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), FONT)
        fonts.set(qn("w:hAnsi"), FONT)
        rpr.append(fonts)
        lvl.extend([start, num_fmt, lvl_text, jc, ppr, rpr])
        abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        ref = OxmlElement("w:abstractNumId")
        ref.set(qn("w:val"), str(abs_id))
        num.append(ref)
        numbering.append(num)

    make("bullet", "•", next_abs, next_num)
    make("decimal", "%1.", next_abs + 1, next_num + 1)
    return next_num, next_num + 1


def apply_num(paragraph, num_id: int):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, nid])


def setup_styles(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb("1A1A1A")
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.10
    pf.widow_control = True

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        st = doc.styles[style_name]
        st.font.name = FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = rgb(color)
        st._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.keep_together = True

    caption = doc.styles["Caption"]
    caption.font.name = FONT
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)
    caption._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)


def set_header_footer(section, first=False):
    section.different_first_page_header_footer = first
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    set_run(p.add_run("EMOTIONAL LIFE SUPPORT"), size=8.5, color=MUTED, bold=True)
    p.add_run("\t")
    set_run(p.add_run("Alternatives report | August 2026"), size=8.5, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    set_run(fp.add_run("Page "), size=9, color=MUTED)
    add_field(fp, "PAGE")


def add_title_cover(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(92)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("MARKET & COMPETITOR RESEARCH"), size=10, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_run(p.add_run("Alternatives to\nEmotional Life Support"), size=30, color=INK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    set_run(p.add_run("UK competitive landscape, public pricing and launch implications"), size=15, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(72)
    set_run(p.add_run("Scope: adults rebuilding after major life disruption, with separation and divorce as the provisional beachhead"), size=10.5, color=MUTED, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    set_run(p.add_run("Prepared for Chris"), size=11, color=INK, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("Research snapshot: 15 August 2026"), size=10, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(44)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Prices are public list prices and may change. Provider claims are not independent outcome evidence."), size=8.5, color=MUTED, italic=True)


def add_paragraph(doc, text="", *, bold_lead=None, italic=False, color=None, after=None, before=None, align=None):
    p = doc.add_paragraph()
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    keep_lines(p)
    if bold_lead and text.startswith(bold_lead):
        set_run(p.add_run(bold_lead), bold=True, color=color)
        set_run(p.add_run(text[len(bold_lead):]), italic=italic, color=color)
    else:
        set_run(p.add_run(text), italic=italic, color=color)
    return p


def add_bullet(doc, text: str, bullet_id: int, *, bold_lead=None, after=4):
    p = doc.add_paragraph()
    apply_num(p, bullet_id)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.167
    keep_lines(p)
    if bold_lead and text.startswith(bold_lead):
        set_run(p.add_run(bold_lead), bold=True)
        set_run(p.add_run(text[len(bold_lead):]))
    else:
        set_run(p.add_run(text))
    return p


def add_number(doc, text: str, number_id: int, *, bold_lead=None):
    p = doc.add_paragraph()
    apply_num(p, number_id)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.167
    keep_lines(p)
    if bold_lead and text.startswith(bold_lead):
        set_run(p.add_run(bold_lead), bold=True)
        set_run(p.add_run(text[len(bold_lead):]))
    else:
        set_run(p.add_run(text))
    return p


def add_callout(doc, label: str, text: str, *, fill="EAF2F8", border=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.10
    set_paragraph_shading(p, fill, border)
    set_run(p.add_run(label.upper() + "  "), size=9.5, color=border, bold=True)
    set_run(p.add_run(text), size=10.5, color=INK)
    keep_lines(p)
    return p


def add_heading(doc, text: str, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_keep_with_next(p)
    return p


def add_caption(doc, text: str):
    p = doc.add_paragraph(text, style="Caption")
    set_keep_with_next(p)
    return p


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[int], *, font_size=9.2, header_fill=LIGHT_GRAY, first_col_bold=False, aligns: Sequence[str] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, label in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if aligns and aligns[i] == "center" else WD_ALIGN_PARAGRAPH.LEFT
        set_run(p.add_run(label), size=font_size, color=INK, bold=True)
    for ridx, row_values in enumerate(rows):
        row = table.add_row()
        if ridx % 2 == 1:
            for c in row.cells:
                set_cell_shading(c, VERY_LIGHT)
        for i, value in enumerate(row_values):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if aligns and aligns[i] == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run(p.add_run(str(value)), size=font_size, color="1A1A1A", bold=(first_col_bold and i == 0))
    set_table_geometry(table, widths)
    return table


def add_pros_cons(doc, pros: Sequence[str], cons: Sequence[str], bullet_id: int):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run("Pros"), size=10.5, color=GREEN, bold=True)
    for item in pros:
        add_bullet(doc, item, bullet_id, after=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run("Cons / limits"), size=10.5, color=RED, bold=True)
    for item in cons:
        add_bullet(doc, item, bullet_id, after=2)


def page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()
    setup_styles(doc)
    bullet_id, number_id = ensure_custom_numbering(doc)
    set_header_footer(doc.sections[0], first=True)
    doc.core_properties.title = "Alternatives to Emotional Life Support"
    doc.core_properties.subject = "UK competitive landscape, public pricing and launch implications"
    doc.core_properties.author = "OpenAI Codex for Chris"
    doc.core_properties.keywords = "Emotional Life Support, competitors, alternatives, divorce coaching, pricing, UK"
    doc.core_properties.comments = "Research snapshot dated 15 August 2026"

    add_title_cover(doc)
    page_break(doc)

    add_heading(doc, "Executive summary", 1)
    add_callout(
        doc,
        "Bottom line",
        "There is a credible gap for one trusted human who combines emotional steadiness, transition clarity, practical organisation and follow-through. The gap is not empty, however: customers can assemble much of the value from a therapist, a divorce or life coach, friends, a virtual assistant and specialist advisers. Emotional Life Support must therefore sell the integration and navigation—not claim to replace regulated expertise.",
    )
    conclusions = [
        "The closest direct alternative in a separation/divorce beachhead is a divorce coach. Public examples range from about £58–£160 per one-hour session, with ongoing high-access packages at roughly £375–£800 per month. [9–12]",
        "Private counselling is the strongest substitute for emotional depth and safety. More than half of BACP-directory listings advertise £40–£60 per session, while Counselling Directory reported a January 2026 database average of £60.60—about £263 per month at weekly frequency. [3–4]",
        "Low-cost substitutes are abundant: friends and family (£0), NHS Talking Therapies (£0 for eligible users), crisis listening (£0), online communities, £20-per-month memberships and AI/self-help tools. [2,17,25–28]",
        "Specialists solve narrower but urgent jobs. A divorce application in England and Wales costs £628; a MIAM is usually around £120; and comprehensive amicable legal/negotiation services list from £1,200 to £5,910. [19–21]",
        "£1,000 per month is not impossible, but it is a premium price. It sits above common weekly coaching and therapy, near the upper end of published divorce-coaching retainers, and below a small ultra-premium life-coaching market. It needs visibly more than four conversations.",
        "The biggest commercial risk is category confusion. Life coaching is not regulated in the UK, while the proposed service sits close to therapy, crisis support, legal guidance and domestic-abuse safeguarding. Trust architecture is part of the product, not compliance copy added later. [5–6]",
    ]
    for x in conclusions:
        add_bullet(doc, x, bullet_id)

    add_caption(doc, "Table 1. Price anchors a prospective customer may encounter (public list prices, observed 15 August 2026)")
    add_table(
        doc,
        ["Alternative", "Public price anchor", "What that usually buys", "Competitive implication"],
        [
            ["Free / informal", "£0", "Friends, family, NHS, charities, helplines, forums", "Always the default competitor"],
            ["Private therapy", "Often £40–£60/session; £60.60 database average", "Usually one 50–60 minute session; weekly ≈ £173–£263/month", "Strong emotional and professional trust benchmark"],
            ["Divorce coaching", "≈ £58–£160/session; £375–£800/month", "1:1 support, planning and sometimes messaging", "Closest direct comparator for the beachhead"],
            ["Groups / courses", "£20/month to £1,099 programme", "Community, curriculum, group calls; some 1:1 upgrades", "Far cheaper scalability and belonging"],
            ["Accountability / PA", "£129–£590/month; VA £40/hour + VAT", "Check-ins, execution, reminders or delegated admin", "Strong on follow-through, weak on emotional depth"],
            ["Legal / mediation", "MIAM ≈ £120; services £1,200–£5,910", "Regulated or accredited problem-solving", "High urgency and authority, but narrow scope"],
        ],
        [1700, 1950, 2880, 2830],
        font_size=8.8,
        first_col_bold=True,
    )

    add_heading(doc, "Recommendation in one sentence", 2)
    add_paragraph(
        doc,
        "Test Emotional Life Support as a bounded, high-trust ‘life rebuild navigator’ for people who are functional enough for coaching but overwhelmed by the emotional, practical and identity consequences of separation—using a stepped offer around £550–£700 per month and reserving £850–£1,100 for a genuinely intensive tier with explicit practical-support capacity.",
    )

    add_heading(doc, "1. Scope, definitions and method", 1)
    add_heading(doc, "1.1 What is being compared", 2)
    add_paragraph(
        doc,
        "The working concept is a premium, non-clinical, one-to-one service for adults whose old life has been disrupted and who feel lost about what comes next. The intended mechanism combines emotional support, self-understanding, practical sparring, hands-on organisation, accountability and specialist navigation. Separation/divorce is treated as the leading test segment, not a permanent niche. [1]",
    )
    add_paragraph(
        doc,
        "An ‘alternative’ is anything a customer may hire, access or rely on to make progress on the same job. This includes direct competitors, adjacent professionals, free services, informal support and doing nothing. It does not imply equivalence. A therapist, mediator and virtual assistant solve different problems and should not be collapsed into one category.",
    )

    add_heading(doc, "1.2 Evidence labels", 2)
    for item in (
        "Evidence: current public prices, official scope statements, professional-body guidance and published research.",
        "Inference: what those offerings are likely to mean for customer choice and market positioning.",
        "Estimate: normalised monthly costs, job-coverage scores and proposed pricing ranges calculated for this report.",
        "Hypothesis: what Emotional Life Support could offer or charge; none of these hypotheses is validated demand.",
    ):
        add_bullet(doc, item, bullet_id)

    add_heading(doc, "1.3 Research limitations", 2)
    for item in (
        "This is a directional UK desk-research snapshot, not a census of the market. Public-pricing samples overrepresent providers that publish prices.",
        "Provider websites establish what is offered and claimed, not whether outcomes are achieved. Testimonials were not treated as independent proof.",
        "Prices may exclude VAT, court fees, travel, assessments, add-ons or specialist work. One-off packages have not been forced into monthly equivalents when duration is unclear.",
        "The evidence base for generic life coaching remains limited and heterogeneous. A review found only five eligible intervention studies, two randomised, and concluded that only tendencies could be identified. [7]",
        "A 2011 meta-analysis reported a moderate mean effect across separation/divorce interventions, but these were heterogeneous programmes; it does not validate the commercial category of divorce coaching or the proposed service. [8]",
    ):
        add_bullet(doc, item, bullet_id)

    add_callout(
        doc,
        "Important boundary",
        "This report compares purchasing alternatives. It does not recommend substituting Emotional Life Support for therapy, crisis response, safeguarding, legal advice, financial advice or medical care.",
        fill="FFF7E6",
        border=GOLD,
    )

    add_heading(doc, "2. The alternative landscape", 1)
    add_paragraph(doc, "Customers are not choosing between neat categories. They are trying to complete several overlapping jobs, often under time pressure:")
    for item in (
        "Feel heard, less alone and sufficiently steady to think.",
        "Understand what has happened, who they are now and what they want next.",
        "Make decisions and convert them into manageable actions.",
        "Organise documents, appointments, research, routines and competing demands.",
        "Get qualified answers on mental health, law, money, health or safety.",
        "Sustain momentum without becoming dependent on a helper.",
    ):
        add_bullet(doc, item, bullet_id)

    add_caption(doc, "Table 2. Estimated job coverage by service model (1 = weak, 5 = strong). These are analyst estimates based on delivery model, not measured outcomes.")
    add_table(
        doc,
        ["Alternative", "Emotional", "Clarity", "Action", "Practical", "Specialist", "Continuity"],
        [
            ["Self / friends & family", "4", "2", "2", "2", "1", "2"],
            ["NHS / charity / helpline", "3", "2", "1", "1", "4", "1"],
            ["Private therapy", "5", "4", "2", "1", "4", "3"],
            ["Divorce coaching", "4", "4", "4", "3", "2", "4"],
            ["General transition coaching", "3", "4", "5", "2", "1", "4"],
            ["Courses / peer groups", "4", "3", "2", "1", "2", "3"],
            ["Legal / mediation", "1", "3", "3", "3", "5", "2"],
            ["Accountability / PA", "1", "2", "5", "5", "1", "4"],
            ["Self-help / AI", "2", "3", "3", "1", "1", "5"],
            ["ELS intended design*", "5", "5", "5", "4", "2", "5"],
        ],
        [2460, 1050, 1050, 1050, 1050, 1350, 1350],
        font_size=8.7,
        first_col_bold=True,
        aligns=["left", "center", "center", "center", "center", "center", "center"],
    )
    add_paragraph(doc, "* Intended design only. The low specialist-authority score is deliberate: the service should navigate to specialists, not impersonate them.", italic=True, color=MUTED, after=10)

    add_heading(doc, "2.1 Strategic reading of the map", 2)
    for item in (
        "The white space is integration. Most alternatives are strong in one or two jobs and weak elsewhere.",
        "The strongest direct competitors are divorce coaches because they combine emotional support, situational relevance, decision help and messaging.",
        "The strongest trust competitor is therapy. Buyers may ask why they should pay more for a less-regulated service.",
        "The strongest practical competitor is a PA, concierge or accountability coach. These services expose whether ‘hands-on help’ is real capacity or merely supportive conversation.",
        "The strongest price competitor is self-reliance supported by friends, free information and low-cost tools.",
    ):
        add_bullet(doc, item, bullet_id)

    add_heading(doc, "3. Detailed assessment of each alternative", 1)

    add_heading(doc, "3.1 Doing nothing, self-reliance, friends and family", 2)
    add_paragraph(doc, "Description and price: The person delays a purchase, works it out alone, or relies on trusted people. Direct price is £0, although the hidden cost may be time, emotional load, stalled decisions or strain on relationships.")
    add_pros_cons(
        doc,
        ["Immediate, familiar and financially accessible.", "Can provide love, context, practical help and social belonging.", "Preserves autonomy and avoids the risk of a poor professional fit."],
        ["Advice can be biased, inconsistent or entangled with loyalties.", "Loved ones may be emotionally involved, exhausted or unavailable at the exact moments of need.", "Usually lacks structure, confidentiality, specialist navigation and reliable follow-through."],
        bullet_id,
    )
    add_callout(doc, "Competitive implication", "This is the real default. The service must demonstrate why a neutral, structured relationship changes the outcome—not merely offer another sympathetic conversation.")

    add_heading(doc, "3.2 NHS, EAP, charities and crisis listening", 2)
    add_paragraph(doc, "Description and price: NHS Talking Therapies provides evidence-based treatment for specified anxiety and depression presentations and is free for patients; many adults in England can self-refer. Employer EAPs, universities and charities may also provide short-term sessions. Samaritans and Shout provide free listening or crisis support. [2,27–28]")
    add_pros_cons(
        doc,
        ["Free at point of use and embedded in clearer clinical or safeguarding pathways.", "Appropriate route for clinical symptoms, acute distress and crisis support.", "Strong public trust and referral legitimacy."],
        ["Eligibility, waiting time, locality and number of sessions may limit access or continuity.", "Services are designed around mental-health or crisis needs, not whole-life rebuilding, detailed administration or open-ended practical sparring.", "A client may still need legal, financial, career or day-to-day help elsewhere."],
        bullet_id,
    )
    add_callout(doc, "Competitive implication", "Emotional Life Support should be designed to refer into and, where appropriate, complement these services. Competing on ‘support when distressed’ would be unsafe and commercially weak.")

    add_heading(doc, "3.3 Private counselling and psychotherapy", 2)
    add_paragraph(doc, "Description and price: A trained practitioner offers confidential talking therapy, commonly in weekly 50–60 minute sessions. More than half of BACP-directory listings advertise £40–£60 per session; Counselling Directory reported an average listed fee of £60.60 in January 2026. Weekly service therefore commonly normalises to approximately £173–£263 per month, with higher-cost specialists reaching £100–£150+ per session. [3–4]")
    add_pros_cons(
        doc,
        ["Strongest alternative for emotional depth, mental-health formulation, trauma-aware work and professional boundaries.", "Recognisable professional registers, ethics, supervision and complaints routes improve trust when a registered practitioner is chosen.", "Can address patterns and distress that coaching should not treat."],
        ["Not normally designed to organise the client’s practical life, coordinate specialists or complete tasks.", "Between-session access and rapid tactical decision support are often limited.", "The buyer may experience progress as less concrete if goals and outcomes are not explicit; fit varies by practitioner and modality."],
        bullet_id,
    )
    add_callout(doc, "Competitive implication", "A £1,000 service costs roughly four times the current weekly-therapy price anchor. It needs a different job, more access and tangible practical value—not a claim that it is deeper or better than therapy.")

    add_heading(doc, "3.4 Divorce coaching", 2)
    add_paragraph(doc, "Description and price: One-to-one support aimed at the emotional, communication, planning and practical demands of separation. Public examples span approximately £58–£160 per 60-minute session. Ongoing packages include £375 per month with four sessions and WhatsApp support, £520 per month with four sessions and limited text/call access, and £800 per month with four sessions, unlimited texts and one emergency call per week. [9–12]")
    add_pros_cons(
        doc,
        ["Highly recognisable trigger, language and moment of urgency.", "Often combines reassurance, planning, communication support, accountability and between-session access.", "Can reduce the coordination burden around legal meetings and paperwork without charging solicitor rates."],
        ["Life coaching is unregulated in the UK; training, scope, complaints and safeguarding vary. [5]", "Some offers blur emotional support, legal-form assistance and ‘emergency’ contact without making competence boundaries obvious.", "Most public proof is practitioner-selected testimonials; independent evidence for commercial divorce coaching specifically is weak."],
        bullet_id,
    )
    add_callout(doc, "Competitive implication", "This is the closest benchmark. Emotional Life Support must either be more integrated and hands-on, or charge nearer the £375–£600 market centre while it establishes trust and evidence.")

    add_heading(doc, "3.5 General life, transition and premium personal coaching", 2)
    add_paragraph(doc, "Description and price: Coaching helps clients clarify values, decisions and goals, then act. A public transition package costs £949 for one two-hour and three 90-minute sessions. Accountability Coaching London lists £590 per month for weekly calls and between-session checkpoints. At the premium end, Nick Hatter lists £4,320 for six weeks and £7,200 for three months. [15–16,24]")
    add_pros_cons(
        doc,
        ["Strong fit for future vision, identity, decision-making, confidence and action.", "Commercial norms support packages, assessments, homework, messaging and outcomes-oriented work.", "Premium examples prove that high prices exist, though not that this audience will pay them."],
        ["Generic ‘transformation’ language is crowded and may feel tone-deaf during acute disruption.", "Quality and safeguarding vary because coaching is not regulated; an ICF credential can be a trust signal but is not statutory regulation. [5–6]", "Premium pricing often rests on founder reputation, affluent audiences and social proof that a new service will not initially possess."],
        bullet_id,
    )
    add_callout(doc, "Competitive implication", "Use coaching methods internally, but acquire customers through the event and job they recognise. ‘Life coaching’ alone is unlikely to communicate the required safety, specificity or practical help.")

    add_heading(doc, "3.6 Courses, workshops, memberships and peer groups", 2)
    add_paragraph(doc, "Description and price: Scalable programmes combine curriculum, exercises and peer connection. Examples include amicable space at £20 per month; Divorce Recovery Workshop residential weekends at £190–£280; an online course at £249, a 12-week group version at £699, and hybrid upgrades at £899–£1,099. [13,17–18]")
    add_pros_cons(
        doc,
        ["Normalises the experience and reduces isolation through people who ‘get it’.", "Much lower cost per person and a clear curriculum can make progress feel structured.", "Useful future scale route for Emotional Life Support once repeatable elements are discovered."],
        ["Group timing, disclosure and dynamics do not suit everyone; risk and abuse contexts require careful screening.", "Advice from peers can be inaccurate or polarising.", "Less responsive to a person’s exact decisions, practical workload or changing intensity."],
        bullet_id,
    )
    add_callout(doc, "Competitive implication", "Do not build a course first. Use early one-to-one work to identify repeatable tools, then offer peer support as a deliberate complement rather than a cheaper imitation of personal support.")

    add_heading(doc, "3.7 Solicitors, mediators, financial advisers and other specialists", 2)
    add_paragraph(doc, "Description and price: These providers answer narrow, consequential questions with recognised competence. In England and Wales the court fee to apply for divorce is £628. A MIAM usually costs around £120; eligible families may receive up to £500 toward mediation. amicable lists £1,200–£1,800 for automated or assisted divorce/consent-order services and £3,450–£5,910 for negotiation services. [19–21]")
    add_pros_cons(
        doc,
        ["Highest authority for legal, financial, medical or safeguarding questions.", "Can create binding documents, regulated advice or accredited mediation outcomes.", "High willingness to pay when the consequence of error is material."],
        ["Narrow remit and professional time can be expensive.", "The client still has to organise facts, prepare questions, cope emotionally and implement decisions.", "A specialist may not provide continuity across the person’s whole transition."],
        bullet_id,
    )
    add_callout(doc, "Competitive implication", "The opportunity is to make specialist time work better: organise the client, prepare questions, track decisions and translate next steps—while never interpreting the law or giving regulated advice.")

    add_heading(doc, "3.8 Accountability coaching, virtual assistants and concierge support", 2)
    add_paragraph(doc, "Description and price: Accountability services create daily or weekly follow-through; public examples run from £129–£229 per month for WhatsApp-based coaching to £590 per month for weekly calls and messaging. A VA example lists £40 per hour, £185 for five hours and £740 for 20 hours, plus VAT. [22–24]")
    add_pros_cons(
        doc,
        ["Strongest alternative for execution, reminders, research, diary management and reducing cognitive load.", "Clear unit of value—hours, completed tasks or check-ins—makes pricing easier to understand.", "Can be flexed up or down with workload."],
        ["Usually weak on grief, identity, self-worth, emotional regulation and sensitive judgment.", "Delegation can undermine agency if tasks are simply taken over rather than used to rebuild capability.", "Handling deeply personal data across multiple systems creates confidentiality risk."],
        bullet_id,
    )
    add_callout(doc, "Competitive implication", "If Emotional Life Support promises hands-on organisation, it should state the included capacity. Otherwise buyers can reasonably compare it with £185 of concrete PA time or a £590 accountability retainer.")

    add_heading(doc, "3.9 Self-help, apps and AI wellbeing tools", 2)
    add_paragraph(doc, "Description and price: Books, podcasts, courses, meditation products and AI tools offer private, on-demand support. Wysa has a free tier and the UK App Store listed Premium at £11.99 monthly or £66.99 annually, alongside other coaching/therapist purchases. Wysa explicitly states that the AI is not a replacement for qualified care and is not suitable for crisis or severe and enduring mental illness. [25–26]")
    add_pros_cons(
        doc,
        ["Immediate, low-cost, private and available at any hour.", "Good for psychoeducation, journalling prompts, routines and repeated practice.", "Can scale exercises discovered through high-touch work."],
        ["Limited human judgment, relational accountability and context; confident-sounding errors remain possible.", "Sensitive personal data, retention, provider access and model use require scrutiny.", "Availability can encourage over-reliance while creating an illusion of professional care."],
        bullet_id,
    )
    add_callout(doc, "Competitive implication", "AI should be a bounded support layer or practitioner tool, not the trust proposition. Any client-facing use needs explicit consent, data minimisation, escalation and a non-crisis message.")

    add_heading(doc, "3.10 The assembled support stack", 2)
    add_paragraph(doc, "Description and price: A resourceful client combines several alternatives—for example weekly private therapy, a divorce coach and five to ten hours of VA support. Using current anchors, £60.60 weekly therapy (≈£263/month) + £375 monthly divorce coaching + five VA hours (£185 + VAT) totals about £860/month including 20% VAT on the VA. With ten VA hours, the stack is about £1,082/month. [4,10,22]")
    add_pros_cons(
        doc,
        ["Each professional stays close to their competence and the client can replace one component without losing everything.", "Offers genuine emotional, transition and practical coverage.", "The combined budget shows how a £750–£1,000 integrated service could be economically intelligible."],
        ["The client becomes project manager at the moment they are least able to coordinate.", "Information fragments across providers and advice may conflict.", "Emotional, legal and practical services still cannot be safely collapsed into one unqualified role."],
        bullet_id,
    )
    add_callout(doc, "Competitive implication", "This stack—not a single coach—is the most useful comparison. Emotional Life Support can reduce coordination friction, but must retain a referral network and make clear which specialist costs remain separate.")

    add_heading(doc, "4. Public competitor and price snapshot", 1)
    add_paragraph(doc, "The following prices were visible on public pages on 15 August 2026. They are examples, not averages, and should be rechecked before external use.")

    add_caption(doc, "Table 3. Separation/divorce-specific support")
    add_table(
        doc,
        ["Provider", "Offer", "Public price", "Included / positioning"],
        [
            ["DC² Divorce Coaching", "3 / 6 / 9 sessions", "£180 / £350 / £525", "One-hour sessions; ongoing guidance; ≈£58–£60/hour"],
            ["DC² Divorce Coaching", "Form E support", "£700", "12 hours incl. coaching, document review and liaison"],
            ["Sophie Buck", "Core Monthly", "£375/month", "Four one-hour sessions + WhatsApp"],
            ["Sophie Buck", "Turning Point", "£650 / 6 weeks", "6.5 session hours, WhatsApp, two reassurance calls"],
            ["Sophie Buck", "New Start", "£995 / 3 months", "10.5 session hours, WhatsApp, three reassurance calls"],
            ["D Divorce Coaching", "Single / 3 / 6", "£160 / £460 / £930", "Private one-hour sessions"],
            ["Samantha Johnson", "Bronze", "£520/month", "Four sessions, limited texts, one emergency call"],
            ["Samantha Johnson", "Platinum", "£800/month", "Four sessions, unlimited texts, weekly emergency call"],
            ["Space For Us", "Breaking Up course", "£960 / 16 weeks", "16 weekly 1:1 sessions; £800 upfront option"],
        ],
        [1670, 1680, 1690, 4320],
        font_size=8.5,
        first_col_bold=True,
    )
    add_paragraph(doc, "Sources: [9–12,14].", italic=True, color=MUTED, before=4, after=10)

    add_caption(doc, "Table 4. Courses, memberships and adjacent coaching")
    add_table(
        doc,
        ["Provider", "Offer", "Public price", "Included / positioning"],
        [
            ["amicable space", "Membership", "£20/month", "Content, webinars, forum, 30-minute consultation"],
            ["Divorce Recovery Workshop", "Residential", "£190–£280", "Two-day charity workshop; emotional recovery and peers"],
            ["Emma Rees-Davies", "Self-guided / group", "£249 / £699", "Course alone or course + 12 weekly group sessions"],
            ["Emma Rees-Davies", "Hybrid / VIP", "£899 / £1,099", "Group + 3 individual sessions, or course + 12 sessions"],
            ["Transition Solutions", "Transition package", "£949", "6.5 hours of 1:1 coaching incl. initial consultation"],
            ["Accountability Coaching London", "Full support", "£590/month", "Weekly calls + between-session checkpoints"],
            ["Nick Hatter", "Premium coaching", "£4,320 / 6 weeks", "Weekly sessions, assessment, assignments, 24/7* support"],
        ],
        [1760, 1600, 1750, 4250],
        font_size=8.5,
        first_col_bold=True,
    )
    add_paragraph(doc, "Sources: [13,15–18,24]. *Provider states subject to availability and reasonable use.", italic=True, color=MUTED, before=4, after=10)

    add_caption(doc, "Table 5. Specialist and practical alternatives")
    add_table(
        doc,
        ["Provider / route", "Public price", "What it solves", "Key boundary"],
        [
            ["Private therapist", "Often £40–£60/session; £60.60 average listing", "Emotional/mental-health work", "Not usually admin or specialist coordination"],
            ["GOV.UK divorce application", "£628", "Formal divorce application", "Not financial settlement or personal support"],
            ["MIAM / mediation", "MIAM ≈ £120; voucher up to £500", "Explore and negotiate child/financial arrangements", "Not counselling; agreement may need legal formalisation"],
            ["amicable", "£1,200–£5,910", "Divorce, consent order and negotiation services", "Couple/legal process focus"],
            ["Action PAs", "£40/hour; 5h £185; 20h £740 + VAT", "Delegated admin and operational support", "Not emotional or clinical support"],
            ["Transform Each Day", "£129–£229/month", "Daily WhatsApp accountability", "No scheduled therapy or specialist advice"],
            ["Wysa", "Free; £11.99/month Premium shown", "AI reflection and wellbeing tools", "Not crisis care or replacement for a professional"],
        ],
        [2050, 2100, 2630, 2580],
        font_size=8.5,
        first_col_bold=True,
    )
    add_paragraph(doc, "Sources: [3–4,19–23,25–26].", italic=True, color=MUTED, before=4, after=10)

    add_heading(doc, "5. What £500–£1,000 currently buys", 1)
    add_caption(doc, "Table 6. Like-for-like is impossible; this shows the customer’s visible trade-offs")
    add_table(
        doc,
        ["Budget", "Observable alternatives", "Likely customer conclusion", "Implication for Emotional Life Support"],
        [
            ["≈ £500", "Four weekly divorce-coaching sessions + messaging (£375–£520); or ≈8 weekly therapy sessions at £60.60; or 10 VA hours £370 + VAT", "A meaningful month of human support is available below £500", "A base offer at this level can be credible before strong proof exists"],
            ["≈ £750", "Six-session divorce package (£930, one-off); intensive six-week package £650; 20 VA hours £740 + VAT; group/hybrid programmes £699–£899", "Buyer expects either intensity, many hours or a structured programme", "Define access, response times, practical hours and duration precisely"],
            ["≈ £1,000", "Ten-session divorce programme over three months (£995); 16 weekly guided sessions (£960); 12-session + course VIP (£1,099)", "£1,000 can buy substantial contact time, although often spread over months", "Four standard sessions alone will look thin"],
            ["£1,500+ / month", "Ultra-premium life coaching equivalents; multi-provider stacks; complex legal work", "Reputation, urgency or specialist consequence drives spend", "Not an appropriate launch anchor without trust, evidence and affluent demand"],
        ],
        [1100, 3220, 2510, 2530],
        font_size=8.7,
        first_col_bold=True,
    )

    add_heading(doc, "5.1 The £1,000 test", 2)
    add_paragraph(doc, "At £1,000 per month, four clients produce £4,000 monthly revenue and five produce £5,000 before insurance, supervision, software, taxes, marketing, non-billable administration, cancellations and leave. If the service consumes 10–14 founder hours per client each month, effective revenue is roughly £71–£100 per service hour before overhead. This is commercially plausible but not obviously generous for emotionally demanding, high-liability work.")
    add_callout(
        doc,
        "Pricing conclusion",
        "£1,000 should be tested as an intensive tier, not treated as the single default. The customer should be able to point to concrete capacity—longer sessions, bounded rapid contact, practical organisation and structured reviews—that cannot be obtained from a £375–£520 monthly coach.",
    )

    add_heading(doc, "5.2 Estimated offer ladder to test", 2)
    add_table(
        doc,
        ["Hypothesis", "Indicative price", "Indicative capacity", "Why test it"],
        [
            ["Reset / diagnostic sprint", "£295–£450 once", "Two deep sessions, situation map, priorities and referrals over 2–3 weeks", "Lower-risk way to experience the method; reveals urgent jobs"],
            ["Rebuild", "£550–£700/month", "Four sessions, bounded weekday messaging, action plan and review", "Competitive with strong coaching while allowing broader integration"],
            ["Intensive navigation", "£850–£1,100/month", "Four longer sessions, short reset calls, bounded messaging, 3–4 hours practical work, specialist coordination", "Tests the original premium hypothesis with visible substance"],
            ["Maintenance", "£175–£300/month", "One session, light check-ins and quarterly review", "Supports step-down and reduces dependency incentives"],
        ],
        [1800, 1550, 3670, 2340],
        font_size=8.7,
        first_col_bold=True,
    )
    add_paragraph(doc, "These are research hypotheses, not recommended published prices. They should be tested through paid behaviour with clear eligibility and scope.", italic=True, color=MUTED, before=4, after=10)

    add_heading(doc, "6. Where Emotional Life Support can credibly differentiate", 1)
    add_heading(doc, "6.1 Defensible gaps", 2)
    gaps = [
        ("One relationship across the messy middle.", "A consistent person can connect emotion, meaning, decisions, appointments and action while specialists remain in their lanes."),
        ("Emotional safety plus visible movement.", "The service can make being heard the condition for action, then track whether decisions and practical steps actually happen."),
        ("Variable intensity with deliberate step-down.", "Support can rise around shocks and milestones, then reduce as capacity returns. This turns agency into an operating rule."),
        ("Specialist navigation.", "Preparing questions, organising documents and closing the loop after appointments is valuable without pretending to offer legal, financial, medical or clinical advice."),
        ("Practical organisation as a real deliverable.", "Situation maps, decision logs, next-action plans and bounded research distinguish the service from conversation-only coaching."),
    ]
    for lead, rest in gaps:
        add_bullet(doc, lead + " " + rest, bullet_id, bold_lead=lead)

    add_heading(doc, "6.2 Weak or dangerous differentiation claims", 2)
    for item in (
        "‘Therapy plus action’—this invites clinical comparison and may imply unqualified treatment.",
        "‘Available whenever you need me’—this creates crisis expectations, dependency and unbounded founder capacity.",
        "‘One person for everything’—no individual should be positioned as the source of legal, financial, medical and mental-health authority.",
        "‘Transform your life’—generic, unprovable and mismatched with the buyer’s recognised immediate problem.",
        "A proprietary method before evidence—early structure should be described as a working process, not proven intellectual property.",
    ):
        add_bullet(doc, item, bullet_id)

    add_heading(doc, "6.3 Provisional positioning", 2)
    add_callout(
        doc,
        "Positioning statement",
        "For people who feel lost and overwhelmed after a separation, Emotional Life Support is a non-clinical life-rebuild service that helps them make sense of what is happening, decide what matters next and turn an unmanageable situation into supported, practical steps—while connecting them with qualified specialists where needed.",
        fill="EDF6F1",
        border=GREEN,
    )
    add_paragraph(doc, "Category label to test: ‘life rebuild support’ or ‘separation transition support’. ‘Life coach’ undersells the hands-on integration; ‘emotional support’ alone undersells action; ‘life support’ may imply emergency or clinical care and therefore requires testing.")

    add_heading(doc, "6.4 Trust architecture required before a pilot", 2)
    for item in (
        "Published non-clinical scope, exclusion criteria and a plain-language client agreement.",
        "Safeguarding and crisis procedures, domestic-abuse awareness and vetted referral routes.",
        "Appropriate coach training, supervision/consultation, insurance and a visible code of ethics.",
        "Bounded between-session access: channel, hours, response target and what counts as urgent.",
        "Data-minimised records, explicit consent, retention/deletion rules and separate controls for any AI use.",
        "A scheduled review asking whether the least intensive effective level of support is now appropriate.",
    ):
        add_bullet(doc, item, bullet_id)

    add_heading(doc, "7. Recommended validation plan", 1)
    add_paragraph(doc, "The research supports a focused next experiment, not a final market verdict.")
    steps = [
        ("Interview 12–15 recently separated people.", "Recruit across amicable, high-conflict, child-related and financially complex situations. Ask what they used, spent, trusted, rejected and still had to coordinate."),
        ("Interview 6–8 referrers.", "Include therapists, family solicitors, mediators, domestic-abuse specialists and divorce coaches. Test complementarity, safety and referral criteria."),
        ("Test three concept cards.", "Compare a £350 reset sprint, a £600 monthly rebuild service and a £950 intensive tier. Show exact capacity and boundaries; do not ask only ‘would you pay?’"),
        ("Run 3–5 bounded paid pilots.", "Use explicit eligibility, a four- to six-week review point, outcome measures, adverse-effect/dependency checks and a step-down decision."),
        ("Track the displaced alternative.", "At purchase, ask: ‘What would you have done instead?’ The answer reveals the real competitor and the correct price anchor."),
    ]
    for lead, rest in steps:
        add_number(doc, lead + " " + rest, number_id, bold_lead=lead)

    add_heading(doc, "7.1 Continue, change or stop criteria", 2)
    add_table(
        doc,
        ["Decision", "Evidence threshold"],
        [
            ["Continue", "At least three eligible users pay; most complete agreed actions; clients can name value beyond ‘being listened to’; referrals and boundaries work; founder load is sustainable."],
            ["Change", "Interest is real but the recognised job is narrower, price clusters below the hypothesis, practical work dominates, or the name creates crisis/therapy confusion."],
            ["Stop / redesign", "Demand depends on unbounded availability, clients primarily need therapy or regulated advice, adverse/dependency signals appear, or paid behaviour does not follow positive interview feedback."],
        ],
        [1800, 7560],
        font_size=9.2,
        first_col_bold=True,
    )

    add_heading(doc, "7.2 Recommended immediate decision", 2)
    add_callout(
        doc,
        "Decision",
        "Proceed with divorce/separation as the initial research beachhead. Design the first offer around integration, practical clarity and step-down—not around unlimited access or a universal promise. Treat £600–£700 per month as the central pilot hypothesis and £950–£1,000 as a separately specified intensive tier.",
        fill="EDF6F1",
        border=GREEN,
    )

    add_heading(doc, "Sources and price notes", 1)
    add_paragraph(doc, "All web sources were accessed on 15 August 2026. Provider prices and descriptions are current public listings at the time of access and can change. Bracketed numbers in the report refer to this list.")
    for i, (title, description, url) in enumerate(SOURCES, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.28)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.05
        set_run(p.add_run(f"[{i}] "), size=9, color=INK, bold=True)
        set_run(p.add_run(title + ". "), size=9, color=INK, bold=True)
        set_run(p.add_run(description), size=9, color="333333")
        if url:
            set_run(p.add_run(" "), size=9)
            add_hyperlink(p, "Source", url)

    add_heading(doc, "Calculation notes", 2)
    for item in (
        "Monthly equivalents use 52 weeks ÷ 12 months = 4.33 weeks per month.",
        "Weekly therapy at £40–£60 equals approximately £173–£260 per month; £60.60 equals approximately £263 per month.",
        "The assembled support stack applies 20% VAT to the cited VA price because that provider states prices are subject to VAT; tax treatment varies by provider and customer.",
        "Ultra-premium coaching monthly equivalents are used only as context: £4,320 over six weeks ≈ £3,118 per month; £7,200 over three months = £2,400 per month; £18,000 over 12 months = £1,500 per month.",
        "Job-coverage scores and proposed offer prices are analyst estimates and must not be presented externally as market evidence.",
    ):
        add_bullet(doc, item, bullet_id)

    # Final section-level formatting consistency.
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
